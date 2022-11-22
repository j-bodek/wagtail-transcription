import requests
from docx import Document as docx_document
import tempfile
import io
import datetime
from typing import Type


class ProcessTranscriptionMixin:
    """
    Contains methods that help to clean transcripted audio data and
    create transcription docx file
    """

    def get_transcription(self, transcript_id: str) -> dict:
        # GET TRANSCRIPED FILE
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcript_id}"
        headers = {
            "authorization": self.api_token,
        }
        r = requests.get(endpoint, headers=headers)
        return r.json()

    def phreses_generator(self, words: list) -> list:
        # phrase dict will store info about start of the
        # sentence, text and speaker
        phrase = {}

        for word in words:
            if phrase.get("speaker") != word["speaker"]:
                if phrase:
                    # if phrase is not empty dict
                    yield phrase

                # redifine phrase object
                phrase = {
                    "start": str(datetime.timedelta(milliseconds=word["start"]))[:-4],
                    "speaker": word["speaker"],
                    "text": word["text"] or "",
                }

            elif phrase["speaker"] == word["speaker"]:
                # append word to phrase text
                phrase["text"] += f" {word['text']}" if word["text"] else ""

        if phrase:
            # return last phrase
            yield phrase

    def create_transcript_docx(self, words: list) -> Type[io.BytesIO]:

        document = docx_document()
        with tempfile.NamedTemporaryFile() as tmp:
            for phrase in self.phreses_generator(words):
                start_paragraph = document.add_paragraph()
                run = start_paragraph.add_run(str(phrase.get("start")))
                run.bold = True
                document.add_paragraph(phrase.get("text") + "\n")

            document.save(tmp.name)
            io_output = io.BytesIO(tmp.read())

        return io_output
